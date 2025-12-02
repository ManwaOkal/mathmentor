"""
Processes teacher-uploaded documents (PDF, DOCX, TXT) for RAG.
"""
import os
import asyncio
import httpx
from typing import List, Dict, Any, Optional
from data_processing.chunkers import MathChunker
from data_processing.embeddings import EmbeddingGenerator
from data_processing.vector_store import VectorStore
from lib.supabase_client import get_supabase_client

# Try to import PDF/DOCX libraries, fallback to basic if not available
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2 as pypdf
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentProcessor:
    """Processes teacher-uploaded documents."""
    
    def __init__(self):
        self.chunker = MathChunker(chunk_size=500, chunk_overlap=50)
        self.embedder = EmbeddingGenerator()
        self.vector_store = VectorStore()
        self.supabase = get_supabase_client()
    
    async def process_document(self, document_id: str) -> Dict[str, Any]:
        """
        Process a document: extract text, chunk, generate embeddings, store.
        Only processes if not already processed.
        
        Args:
            document_id: UUID of the document
            
        Returns:
            Processing result with chunk count
        """
        try:
            # Get document record
            doc_result = self.supabase.table('teacher_documents').select('*').eq('document_id', document_id).single().execute()
            
            if not doc_result.data:
                raise ValueError(f"Document {document_id} not found")
            
            document = doc_result.data
            
            # Check if already processed
            if document.get('status') == 'ready':
                # Check if chunks exist
                chunks_check = self.supabase.table('document_chunks').select('chunk_id').eq('document_id', document_id).limit(1).execute()
                if chunks_check.data:
                    return {
                        'success': True,
                        'chunk_count': document.get('metadata', {}).get('chunk_count', 0),
                        'document_id': document_id,
                        'message': 'Document already processed'
                    }
            
            file_url = document['file_url']
            
            # Update status to processing
            self.supabase.table('teacher_documents').update({
                'status': 'processing'
            }).eq('document_id', document_id).execute()
            
            # Extract text from file
            text = await self._extract_text_from_file(file_url, document.get('file_type'))
            
            if not text:
                raise ValueError("Failed to extract text from document")
            
            # Chunk the document
            chunks = self.chunker.chunk_by_difficulty(
                text,
                document.get('difficulty', 'intermediate'),
                document.get('title', 'Document')
            )
            
            # Add document metadata to chunks
            for i, chunk in enumerate(chunks):
                chunk['document_id'] = document_id
                chunk['page_number'] = chunk.get('metadata', {}).get('page_number', 1)
                chunk['chunk_index'] = i
                chunk['metadata'] = {
                    **chunk.get('metadata', {}),
                    'document_title': document['title'],
                    'document_type': document.get('file_type', 'unknown')
                }
            
            # Generate embeddings
            chunks = self.embedder.embed_chunks(chunks)
            
            # Store chunks in database
            chunk_ids = []
            for chunk in chunks:
                chunk_data = {
                    'document_id': document_id,
                    'content': chunk['content'],
                    'embedding': chunk['embedding'],
                    'page_number': chunk.get('page_number', 1),
                    'chunk_index': chunk.get('chunk_index', 0),
                    'metadata': chunk.get('metadata', {})
                }
                
                result = self.supabase.table('document_chunks').insert(chunk_data).execute()
                if result.data:
                    chunk_ids.append(result.data[0]['chunk_id'])
            
            # Update document status
            self.supabase.table('teacher_documents').update({
                'status': 'ready',
                'metadata': {
                    **document.get('metadata', {}),
                    'chunk_count': len(chunk_ids),
                    'processed_at': 'now()'
                }
            }).eq('document_id', document_id).execute()
            
            return {
                'success': True,
                'chunk_count': len(chunk_ids),
                'document_id': document_id
            }
            
        except Exception as e:
            # Update status to failed
            try:
                self.supabase.table('teacher_documents').update({
                    'status': 'failed',
                    'metadata': {
                        **document.get('metadata', {}),
                        'error': str(e)
                    }
                }).eq('document_id', document_id).execute()
            except:
                pass
            
            raise Exception(f"Failed to process document: {str(e)}")
    
    async def _extract_text_from_file(self, file_url: str, file_type: Optional[str]) -> str:
        """
        Extract text from file based on type.
        Downloads file from URL if needed, then extracts text.
        
        Args:
            file_url: URL or path to file
            file_type: MIME type of file
            
        Returns:
            Extracted text
        """
        # Download file if it's a URL
        file_content = None
        if file_url.startswith('http://') or file_url.startswith('https://'):
            async with httpx.AsyncClient() as client:
                response = await client.get(file_url)
                response.raise_for_status()
                file_content = response.content
        elif file_url.startswith('/') or file_url.startswith('./'):
            # Local file
            try:
                with open(file_url, 'rb') as f:
                    file_content = f.read()
            except Exception as e:
                raise ValueError(f"Could not read local file: {e}")
        else:
            raise ValueError(f"Unsupported file path format: {file_url}")
        
        if not file_content:
            raise ValueError("Could not retrieve file content")
        
        # Extract text based on file type
        if file_type == 'text/plain' or file_url.endswith('.txt'):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin-1')
                except:
                    return file_content.decode('utf-8', errors='ignore')
        
        elif file_type == 'application/pdf' or file_url.endswith('.pdf'):
            if not PDF_AVAILABLE:
                raise ValueError("PDF extraction library not installed. Install with: pip install pypdf")
            
            import io
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"\n--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            return "\n".join(text_parts) if text_parts else ""
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_url.endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise ValueError("DOCX extraction library not installed. Install with: pip install python-docx")
            
            import io
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts) if text_parts else ""
        
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _detect_formulas(self, text: str) -> bool:
        """Detect if text contains mathematical formulas."""
        formula_indicators = [
            '=', '^', '_', '\\frac', '\\sqrt', '\\sum',
            '∫', '∂', 'lim', '→', '∈', '∀', '∃', '$'
        ]
        return any(indicator in text for indicator in formula_indicators)

